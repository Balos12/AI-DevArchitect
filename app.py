import os
import streamlit as st
import openai
from docx import Document
from io import BytesIO
from functools import lru_cache
from datetime import datetime

from dotenv import load_dotenv

load_dotenv()
# ----------------------
# Константы и локализация
# ----------------------
LANGUAGES = {"en": "English", "ru": "Русский"}
TRANSLATIONS = {
    "title": {"en": "AI DevArchitect — Intelligent Architecture Assistant",
              "ru": "AI DevArchitect — ИИ-помощник по проектированию архитектуры"},
    "description": {"en": "First, AI suggests project ideas. Select one to get architecture and multilingual alternatives, including dependencies.",
                    "ru": "Сначала ИИ предложит идеи проектов. Затем выберите одну, чтобы получить архитектуру, зависимости и альтернативы на выбранном языке."},
    "generate_ideas": {"en": "Generate ideas", "ru": "Сгенерировать идеи"},
    "select_project": {"en": "Select project", "ru": "Выберите проект"},
    "generate_architecture": {"en": "Generate architecture", "ru": "Сгенерировать архитектуру"},
    "base_architecture": {"en": "Base Architecture & Dependencies", "ru": "Базовая архитектура и зависимости"},
    "show_alternatives": {"en": "Show alternatives", "ru": "Показать альтернативы"},
    "alternatives": {"en": "Alternative Architectures", "ru": "Альтернативные архитектуры"},
    "analyze_antipatterns": {"en": "Analyze antipatterns", "ru": "Анализ антипаттернов"},
    "antipatterns": {"en": "Detected antipatterns", "ru": "Найденные антипаттерны"},
    "complexity_metrics": {"en": "Complexity & Scalability Metrics", "ru": "Метрики сложности и масштабируемости"},
    "export_arch": {"en": "Download Architecture Report", "ru": "Скачать отчёт по архитектуре"},
    "export_full": {"en": "Download Full Report", "ru": "Скачать полный отчёт"},
    "select_model": {"en": "Select model", "ru": "Выберите модель"},
    "temperature": {"en": "Temperature", "ru": "Температура"},
    "max_tokens": {"en": "Max tokens", "ru": "Максимум токенов"},
    "history": {"en": "Session History", "ru": "История сессий"},
    "clear_history": {"en": "Clear history", "ru": "Очистить историю"},
    "download_history": {"en": "Download History Report", "ru": "Скачать историю"},
    "download_all": {"en": "Download All Reports", "ru": "Скачать всё"}
}

@lru_cache(maxsize=None)
def t(key: str) -> str:
    lang = st.session_state.get("lang", "ru")
    return TRANSLATIONS.get(key, {}).get(lang, key)

# ----------------------
# Настройка OpenAI
# ----------------------
openai.api_key = os.getenv("OPENAI_API_KEY")

# ----------------------
# Интерфейс
# ----------------------
st.set_page_config(page_title="AI DevArchitect", layout="wide")
# Sidebar: language, model, params, history controls
st.sidebar.selectbox(
    "Language / Язык",
    options=list(LANGUAGES.keys()),
    format_func=lambda x: LANGUAGES[x],
    key="lang"
)
available_models = ["gpt-4", "gpt-3.5-turbo"]
st.sidebar.selectbox(
    t("select_model"), options=available_models, key="model"
)
st.sidebar.slider(t("temperature"), 0.0, 1.0, 0.7, 0.05, key="temperature")
st.sidebar.slider(t("max_tokens"), 100, 2000, 800, 100, key="max_tokens")
st.sidebar.markdown("---")
if st.sidebar.button(t("clear_history")):
    st.session_state.history = []

# Initialize history
if "history" not in st.session_state:
    st.session_state.history = []

# Main page
st.title(t("title"))
st.write(t("description"))

# Helper for OpenAI calls
def call_openai_system(prompt: str, max_tokens: int = 500) -> str:
    response = openai.ChatCompletion.create(
        model=st.session_state.get("model", "gpt-4"),
        messages=[{"role": "user", "content": prompt}],
        temperature=st.session_state.get("temperature", 0.7),
        max_tokens=st.session_state.get("max_tokens", 800)
    )
    return response.choices[0].message.content.strip()

# Prompts

def generate_ideas() -> list[str]:
    prompt = ("Сгенерируй 5 идей небольших учебных проектов для студентов-программистов. Пример: To-Do App, Чат, Опросник. Формат:\n1. ..."
              if st.session_state["lang"] == "ru"
              else "Generate 5 small educational project ideas: To-Do App, Chat, Survey. Format:\n1. ...")
    text = call_openai_system(prompt, max_tokens=300)
    return [line.split('. ',1)[1] for line in text.splitlines() if '. ' in line]


def generate_architecture(project: str) -> str:
    prompt_ru = f"""
Ты — помощник по архитектуре ПО. Идея: \"{project}\".
Сгенерируй:
1. Сущности и атрибуты
2. Слои приложения
3. Классы и связи
4. Необходимые зависимости
5. Краткое обоснование
"""
    prompt_en = f"""
You are a software architecture assistant. Idea: \"{project}\".
Generate:
1. Entities and attributes
2. Application layers
3. Classes and relations
4. Required dependencies
5. Brief rationale
"""
    prompt = prompt_ru if st.session_state["lang"] == "ru" else prompt_en
    return call_openai_system(prompt, max_tokens=st.session_state["max_tokens"])


def generate_alternatives(project: str, arch: str) -> str:
    opts = ['Monolith','Microservices','Event-driven','CQRS','CQRS+ES','Clean','Hexagonal','Onion']
    bullet = '\n- '.join(opts)
    prompt_ru = f"""
Проект: \"{project}\".
Архитектура и зависимости:
{arch}

Предложи альтернативы:
- {bullet}
"""
    prompt_en = f"""
Project: \"{project}\".
Architecture and dependencies:
{arch}

Suggest alternatives:
- {bullet}
"""
    prompt = prompt_ru if st.session_state["lang"] == "ru" else prompt_en
    return call_openai_system(prompt)


def analyze_antipatterns(project: str, arch: str) -> str:
    prompt = f"Analyze the following architecture for antipatterns (God Object, Anemic Model, etc.) and suggest fixes.\nProject: {project}\nArchitecture:\n{arch}\n"
    return call_openai_system(prompt)


def evaluate_complexity(project: str, arch: str) -> str:
    prompt = f"Evaluate complexity and scalability of the architecture. Estimate modules, bottlenecks, maintainability.\nProject: {project}\nArchitecture:\n{arch}\n"
    return call_openai_system(prompt)

# UI Actions
if st.button(t("generate_ideas")):
    st.session_state.ideas = generate_ideas()

if "ideas" in st.session_state:
    project = st.selectbox(t("select_project"), st.session_state.ideas)
    if st.button(t("generate_architecture")):
        arch = generate_architecture(project)
        st.session_state.arch = arch
        st.session_state.project = project
        st.subheader(t("base_architecture"))
        st.write(arch)
        # save to history
        st.session_state.history.append({
            'time': datetime.now(), 'project': project,
            'architecture': arch, 'alternatives': None,
            'antipatterns': None, 'complexity': None
        })
        # download arch report
        buf = BytesIO(); doc = Document(); doc.add_heading(project, 0); doc.add_heading(t('base_architecture'),1); doc.add_paragraph(arch); doc.save(buf); buf.seek(0)
        st.download_button(t('export_arch'), data=buf, file_name=f"{project}_architecture.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

if "arch" in st.session_state:
    col1, col2 = st.columns(2)
    with col1:
        if st.button(t("show_alternatives")):
            alts = generate_alternatives(st.session_state.project, st.session_state.arch)
            st.session_state.alts = alts
            st.subheader(t("alternatives")); st.write(alts)
            st.session_state.history[-1]['alternatives'] = alts
            # download full report
            buf_full = BytesIO(); doc_full = Document(); doc_full.add_heading(st.session_state.project,0); doc_full.add_heading(t('base_architecture'),1); doc_full.add_paragraph(st.session_state.arch); doc_full.add_heading(t('alternatives'),1); doc_full.add_paragraph(alts); doc_full.save(buf_full); buf_full.seek(0)
            st.download_button(t('export_full'), data=buf_full, file_name=f"{st.session_state.project}_full_report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    with col2:
        if st.button(t("analyze_antipatterns")):
            antip = analyze_antipatterns(st.session_state.project, st.session_state.arch)
            st.subheader(t("antipatterns")); st.write(antip)
            st.session_state.history[-1]['antipatterns'] = antip
        if st.button(t("complexity_metrics")):
            comp = evaluate_complexity(st.session_state.project, st.session_state.arch)
            st.subheader(t("complexity_metrics")); st.write(comp)
            st.session_state.history[-1]['complexity'] = comp

# Show history
if st.session_state.history:
    st.markdown(f"## {t('history')}")
    for entry in reversed(st.session_state.history):
        ts = entry['time'].strftime('%Y-%m-%d %H:%M:%S')
        st.write(f"**{ts}** — {entry['project']}")
        st.write(entry['architecture'])
        if entry['alternatives']: st.write(entry['alternatives'])
        if entry['antipatterns']: st.write(entry['antipatterns'])
        if entry['complexity']: st.write(entry['complexity'])
        st.markdown("---")

    # Download history report
    def create_history_doc() -> BytesIO:
        doc = Document()
        doc.add_heading(t('history'), 0)
        for entry in st.session_state.history:
            ts = entry['time'].strftime('%Y-%m-%d %H:%M:%S')
            doc.add_heading(f"{ts} — {entry['project']}", level=1)
            doc.add_paragraph(entry['architecture'])
            if entry['alternatives']:
                doc.add_paragraph(entry['alternatives'])
            if entry['antipatterns']:
                doc.add_paragraph(entry['antipatterns'])
            if entry['complexity']:
                doc.add_paragraph(entry['complexity'])
        buf = BytesIO(); doc.save(buf); buf.seek(0)
        return buf

    st.download_button(
        label=t('download_history'),
        data=create_history_doc(),
        file_name="session_history.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    # Download all reports: architecture + alternatives etc. merged
    def create_all_reports_doc() -> BytesIO:
        doc = Document()
        doc.add_heading("All Reports", 0)
        for entry in st.session_state.history:
            ts = entry['time'].strftime('%Y-%m-%d %H:%M:%S')
            doc.add_heading(f"{ts} — {entry['project']}", level=1)
            doc.add_heading(t('base_architecture'), level=2)
            doc.add_paragraph(entry['architecture'])
            if entry['alternatives']:
                doc.add_heading(t('alternatives'), level=2)
                doc.add_paragraph(entry['alternatives'])
            if entry['antipatterns']:
                doc.add_heading(t('antipatterns'), level=2)
                doc.add_paragraph(entry['antipatterns'])
            if entry['complexity']:
                doc.add_heading(t('complexity_metrics'), level=2)
                doc.add_paragraph(entry['complexity'])
        buf = BytesIO(); doc.save(buf); buf.seek(0)
        return buf

    st.download_button(
        label=t('download_all'),
        data=create_all_reports_doc(),
        file_name="all_reports.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
